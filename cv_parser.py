"""
Extracts plain text from a CV file (PDF or DOCX).
No LLM calls here — structured analysis happens in Claude's reasoning,
not in this server, to avoid any extra API cost.
"""
from pathlib import Path
from pypdf import PdfReader
from docx import Document


def extract_text(file_path: str) -> str:
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"CV file not found: {file_path}")

    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf(path)
    elif suffix in (".docx",):
        return _extract_docx(path)
    elif suffix in (".txt", ".md"):
        return path.read_text(errors="ignore")
    else:
        raise ValueError(f"Unsupported CV file type: {suffix}. Use PDF, DOCX, TXT, or MD.")


def _extract_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    return "\n".join(page.extract_text() or "" for page in reader.pages).strip()


def _extract_docx(path: Path) -> str:
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts).strip()
